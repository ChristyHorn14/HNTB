from pathlib import Path

import glog
import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor

from hntb.config_options import HNTBConfig


# Function to convert hexadecimal color code to RGB color
def hex_to_rgb(hex_color):
    return RGBColor(*tuple(int(hex_color[i : i + 2], 16) for i in (0, 2, 4)))


# Function to generate a Word document for each unique attending
def generate_doc_for_attending(output_directory: Path, sorted_df, selected_columns: list, attending):
    # Filter the data for the current attending
    attending_df = sorted_df[sorted_df["Attending"] == attending]

    # Convert attending to string (handle NaN or numeric values)
    attending = str(attending)

    # Create a new Document
    doc = Document()

    # Set document margins (inches)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.orientation = WD_ORIENTATION.LANDSCAPE  # Set orientation to landscape

    # Set default font for the document
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(9)

    # Add a title to the document
    doc.add_heading(f"Attending: {attending}", level=1)
    doc.add_paragraph()  # Add a blank paragraph

    # Add a table to the document
    table = doc.add_table(rows=1, cols=len(selected_columns))

    # Set table style and properties
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER  # Center align table

    # Set column widths based on page margins and number of columns
    total_width_inches = 11  # Landscape paper width in inches
    column_width = total_width_inches / len(selected_columns)

    for cell in table.columns[0].cells:
        cell.width = Inches(column_width)

    # Add header row
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(selected_columns):
        hdr_cells[i].text = col
        hdr_cells[i].paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER  # Center align header text
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER  # Center align vertically
        hdr_cells[i].paragraphs[0].paragraph_format.space_after = Pt(6)  # Add some spacing after header text
        # Set header row font to bold
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    # Add rows from dataframe
    for index, row in attending_df.iterrows():
        row_cells = table.add_row().cells
        for i, col in enumerate(selected_columns):
            if col == "DOB" and pd.notnull(row[col]):
                # Convert datetime to month day year format
                dob_formatted = row[col].strftime("%m/%d/%Y")
                cell_text = dob_formatted
            else:
                cell_text = str(row[col])

            # Check if the column is 'Other Notes' and there's a font_color specified
            if col == "Other Notes" and "font_color" in row.index and pd.notnull(row["font_color"]):
                font_color = hex_to_rgb(row["font_color"])
                run = row_cells[i].paragraphs[0].add_run(cell_text)
                font = run.font
                font.color.rgb = font_color
            else:
                run = row_cells[i].paragraphs[0].add_run(cell_text)

            run.font.name = "Arial"
            run.font.size = Pt(9)
            run.bold = False  # Assuming data should not be bold by default

            # Center align cell text
            row_cells[i].paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER
            row_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            row_cells[i].paragraphs[0].paragraph_format.space_after = Pt(6)  # Add some spacing after cell text

    # Save the document
    doc_path = output_directory / f"PrelimEmail_{attending}.docx"
    doc.save(doc_path)
    glog.info(f"=> Document saved successfully to {doc_path}")


def generate_prelim_emails(cfg: HNTBConfig):
    # Create the directory
    output_directory = cfg.output_directory

    # Read the Excel file from OneDrive
    input_file = cfg.active_tumor_board_file
    sheet_name = "Master Linked"
    df = pd.read_excel(input_file, sheet_name=sheet_name)

    glog.debug(f"=> input_file={input_file}")
    glog.debug(f"=> output_directory={output_directory}")

    # Sort and filter data
    sorted_df = df.sort_values(by=["List", "Sorting Date"], ascending=[True, True])

    # Select only the specified columns
    selected_columns = ["Last name", "First name", "DOB", "Attending", "Summary", "List", "Other Notes"]
    # Get unique values in the "Attending" column
    unique_attending_values = sorted_df["Attending"].unique()

    # Generate Word documents for each unique attending value
    for attending in unique_attending_values:
        generate_doc_for_attending(output_directory, sorted_df, selected_columns, attending)
