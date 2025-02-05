import os
import pandas as pd
from docx import Document
from docxcompose.composer import Composer
from docx.shared import Inches

# Define paths
#input_file = 'Active Tumor Board LINKED.xlsx'
#template_path = '/Users/chrishornung/Desktop/HNTB Reboot/Templates/FacesheetTemplate.docx'
#output_directory = '/Users/chrishornung/Desktop/HNTB Reboot/Outputs'
#output_file = os.path.join(output_directory, 'docs_facesheets.docx')
#header_image_path = '/Users/chrishornung/Desktop/HNTB Reboot/Templates/EVMSLogo.png'  # Update this to your header image path

input_file = './tests/artifacts/hntb_dummy.xlsx'
template_path = './Templates/FacesheetTemplate.docx'
output_directory = './tests/artifacts/Outputs'
output_file = os.path.join(output_directory, 'docs_facesheets.docx')
header_image_path = './Templates/EVMSLogo.png'  # Update this to your header image path

# Load the Excel file and specific sheet
#df = pd.read_excel(input_file, sheet_name='Master Linked')
df = pd.read_excel(input_file)

# Filter out rows where 'List' is '5. Pending'
df = df[df['List'] != '5. Pending']

# Drop blank rows
df.dropna(how='all', inplace=True)

df = df.sort_values(by=["Resident"], ascending=[True])

# Function to check if a value is a date and format it
def format_value(value, not_date=False):
    if pd.isna(value):
        return ""
    if not_date:
        if isinstance(value, float):
            value = int(value)
        return str(value)
    try:
        value = pd.to_datetime(value)
        return value.strftime('%m/%d/%Y')
    except (ValueError, TypeError):
        if isinstance(value, float) and value.is_integer():
            return str(int(value))
        return str(value)


# Function to replace placeholders in the document
def replace_placeholders(document, mapping):
    for paragraph in document.paragraphs:
        for key, value in mapping.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, value)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in mapping.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, value)

# Function to add header to a document
def add_header(document, header_image_path):
    section = document.sections[0]
    header = section.header
    paragraph = header.paragraphs[0]
    run = paragraph.add_run()
    run.add_picture(header_image_path, width=Inches(6))  # Adjust the width as needed

# Function to create a new document with patient data
def create_patient_doc(template_path, patient_data, header_image_path):
    # Load the template document for each patient
    doc = Document(template_path)
    
    # Add header to the document
    add_header(doc, header_image_path)
    
    # Mapping for placeholders and their corresponding data
    mapping = {
        '{First name}': format_value(patient_data['First name']),
        '{Last name}': format_value(patient_data['Last name']),
        '{EPIC MRN}': format_value(patient_data['EPIC MRN'], not_date=True),
        '{DOB}': format_value(patient_data['DOB']),
        '{List}': format_value(patient_data['List']),
        '{Demographics}': format_value(patient_data['Demographics']),
        '{Resident}': format_value(patient_data['Resident']),
        '{Attending}': format_value(patient_data['Attending']),
        '{HP/Clinic}': format_value(patient_data['HP/Clinic']),
        '{Diagnosis}': format_value(patient_data['Diagnosis']),
        '{Imaging1}': format_value(patient_data['Imaging1']),
        '{Imaging 2}': format_value(patient_data['Imaging 2']),
        '{Imaging 3}': format_value(patient_data['Imaging 3']),
        '{OR1}': format_value(patient_data['OR1']),
        '{OR2}': format_value(patient_data['OR2']),
        '{OR3}': format_value(patient_data['OR3']),
        '{Path1}': format_value(patient_data['Path1']),
        '{Path2}': format_value(patient_data['Path2']),
        '{Path3}': format_value(patient_data['Path3']),
        '{Summary}': format_value(patient_data['Summary']),
        '{Other Notes}': format_value(patient_data['Other Notes'])
    }
    
    # Replace placeholders in the document
    replace_placeholders(doc, mapping)
    
    return doc

# Create a master document to collect all patient docs
master_doc = Document()
composer = Composer(master_doc)

# Add header to the master document
add_header(master_doc, header_image_path)

# Iterate over each row in the dataframe and create patient docs
for index, row in df.iterrows():
    patient_doc = create_patient_doc(template_path, row, header_image_path)
    
    # Save the patient document temporarily
    temp_doc_path = f'temp_patient_{index + 1}.docx'
    patient_doc.save(temp_doc_path)
    
    # Append the patient document to the master document
    composer.append(Document(temp_doc_path))
    
    # Delete the temporary file
    os.remove(temp_doc_path)

# Save the master document
composer.save(output_file)
print("Document saved successfully to", output_file)
