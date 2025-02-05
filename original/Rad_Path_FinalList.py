import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Define the file path and sheet name
# file_path = '/Users/chrishornung/Desktop/HNTB Reboot/Active Tumor Board LINKED.xlsx'
# sheet_name = 'Master Linked'
file_path = '/Users/courtneythomas/Library/CloudStorage/OneDrive-OldDominionUniversity/HNTB/Active Tumor Board LINKED.xlsx'
sheet_name = 'Master Linked'

# Read the Excel file
df = pd.read_excel(file_path, sheet_name=sheet_name)


# Filter for "Path"
# path_df = df[(~df['List'].str.contains('Pending')) & (df['Other Notes'].str.contains('PATH', na=False))]
path_df = df[
    (df['List'].str.contains('New patient', na=False))
    | (df['List'].str.contains('Endocrine', na=False))
    | (df['List'].str.contains('Path follow up', na=False))
    ]

# Filter for "RAD"

# TODO: This line will always return an empty df because it is case sensitive
# and 'RAD' never occurs as all caps
rad_df = df[(~df['List'].str.contains('Pending', na=True)) & (df['Other Notes'].str.contains('RAD', na=False))]

# Function to create a table in a Word document
def create_table_in_word(df, output_file):
    document = Document()

    # Add a title to the document
    title = document.add_heading(level=1)
    title_run = title.add_run("Table Export")
    title_run.bold = True
    title_run.font.size = Pt(14)

    # Add a table to the document
    table = document.add_table(rows=1, cols=len(df.columns))
    table.style = 'Table Grid'  # Apply grid style to the table

    # Add header row
    header_row = table.rows[0].cells
    for idx, col in enumerate(df.columns):
        header_row[idx].text = col

    # Add data rows
    for _, row in df.iterrows():
        r = [str(row[col]) for col in df.columns]
        print(r)
        # TODO: This line generates an error, did it ever work?
        table.add_row([str(row[col]) for col in df.columns])
        exit()

    # Adjust table alignment
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Save the document
    document.save(output_file)

# Output file paths
# path_output_file = '/Users/chrishornung/Desktop/HNTB Reboot/Outputs/Path_FinalList.docx'
# rad_output_file = '/Users/chrishornung/Desktop/HNTB Reboot/Outputs/RAD_FinalList.docx'
path_output_file = '/Users/courtneythomas/Library/Mobile Documents/com~apple~CloudDocs/Residency/HNTB/HNTB_Chief/Outputs/Path_FinalList.docx'
rad_output_file = '/Users/courtneythomas/Library/Mobile Documents/com~apple~CloudDocs/Residency/HNTB/HNTB_Chief/Outputs/RAD_FinalList.docx'

# Create tables in Word documents for both sub-dataframes
create_table_in_word(path_df, path_output_file)
create_table_in_word(rad_df, rad_output_file)

print(f"Path sub-data saved to {path_output_file}")
print(f"RAD sub-data saved to {rad_output_file}")
