#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Sat Jul 13 09:14:27 2024

@author: chrishornung
"""
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os
from datetime import datetime


# Read the Excel file
# file_path = '/Users/chrishornung/Desktop/HNTB Reboot/Active Tumor Board LINKED.xlsx'
# sheet_name = 'Master Linked'
# df = pd.read_excel(file_path, sheet_name=sheet_name)
# file_path = '/Users/courtneythomas/Library/CloudStorage/OneDrive-OldDominionUniversity/HNTB/Active Tumor Board LINKED.xlsx'
file_path = "/Users/courtneythomas/Library/Mobile Documents/com~apple~CloudDocs/Residency/HNTB/HNTB Reboot/tests/artifacts/hntb_dummy.xlsx"
sheet_name = 'Master Linked'
df = pd.read_excel(file_path, sheet_name=sheet_name)

# Sort and filter data
sorted_df = df.sort_values(by=["Sorting Date","List"], ascending=[True,True])
selected_columns = ['Last name', 'First name', 'Attending', 'DOB', 'EPIC MRN', 'Diagnosis', 'List', 'Sorting Date']
filtered_df = sorted_df[selected_columns].rename(columns={'EPIC MRN': 'MRN'})

# Insert placeholder columns
filtered_df['Packet'] = ''
filtered_df['Referred by'] = ''
filtered_df['SSN'] = ''
filtered_df['Recommendations'] = ''

# Reorder columns
final_df = filtered_df[['Packet', 'Last name', 'First name', 'Attending', 'SSN', 'DOB', 'MRN', 'Referred by', 'Diagnosis', 'List', 'Recommendations','Sorting Date']]

# Sort final_df by 'List' column in ascending order
final_df = final_df.sort_values(by=['List','Sorting Date'], ascending=[True,True])

desired_lists = ["1. New patient", "2. Endocrine", "3. Path follow up", "4. Radiology follow up"]

# Filter final_df
final_df = final_df[final_df['List'].isin(desired_lists)]
final_df = final_df.drop(columns=['Sorting Date'])


# Function to format DOB and MRN, create Word document with table fitting within page margins
def create_word_document(dataframe, filename):
    document = Document()
    
    # Add table with gridlines
    table = document.add_table(rows=1, cols=len(dataframe.columns))
    table.style = 'Table Grid'
    table.autofit = False
    
    # Set column widths
    section = document.sections[0]
    page_width = section.page_width - section.left_margin - section.right_margin
    column_width = int(page_width / len(dataframe.columns))  # Convert to int
    
    for col in table.columns:
        col.width = column_width
    
    # Add header row
    hdr_cells = table.rows[0].cells
    for i, column in enumerate(dataframe.columns):
        hdr_cells[i].text = column
        hdr_cells[i].paragraphs[0].runs[0].bold = True  # Bold the header cell text
    
    # Add data rows
    for _, row in dataframe.iterrows():
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            if dataframe.columns[i] == 'MRN':
                # Convert MRN to integer, handle NaN
                value = int(value) if pd.notnull(value) else ''
            elif dataframe.columns[i] == 'DOB':
                # Format DOB as month, day, year (without time), handle NaT
                if pd.isnull(value):  # Check for NaT (Not-a-Time)
                    value = ''
                else:
                    value = value.strftime('%B %d, %Y') if isinstance(value, datetime) else ''
            row_cells[i].text = str(value)

    
    # Save document
    document.save(filename)

# Generate document for all attendings
# filename = "/Users/chrishornung/Desktop/HNTB Reboot/Outputs/FinalLists.docx"
# create_word_document(final_df, filename)
# filename = "/Users/courtneythomas/Library/Mobile Documents/com~apple~CloudDocs/Residency/HNTB/HNTB_Chief/Outputs/FinalLists.docx"
filename = "/Users/courtneythomas/Library/Mobile Documents/com~apple~CloudDocs/Residency/HNTB/HNTB Reboot/tests/artifacts/Outputs/FinalLists.docx"
create_word_document(final_df, filename)

print("Word document generated successfully.")
